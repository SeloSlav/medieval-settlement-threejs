from __future__ import annotations

import argparse
import hashlib
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


INLINE_ITALIC = re.compile(r"(<i>.*?</i>)")


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest().upper()


def set_run_font(run, size: float | None = None):
    if size is not None:
        run.font.size = Pt(size)
    if run.font.name:
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:ascii"), run.font.name)
        rfonts.set(qn("w:hAnsi"), run.font.name)


def parse_blocks(path: Path) -> list[str]:
    raw = path.read_text(encoding="utf-8").strip()
    return [block.strip() for block in re.split(r"\r?\n\s*\r?\n", raw) if block.strip()]


def add_inline_runs(paragraph, text: str):
    for part in INLINE_ITALIC.split(text):
        if not part:
            continue
        if part.startswith("<i>") and part.endswith("</i>"):
            run = paragraph.add_run(part[3:-4])
            run.italic = True
        else:
            paragraph.add_run(part)


def add_chapter(document: Document, chapter_file: Path):
    existing = {paragraph.text.strip() for paragraph in document.paragraphs}
    expected_heading = "CHAPTER THREE\nTHE ORCHESTRATORS"
    if expected_heading in existing:
        raise RuntimeError("Chapter Three already exists in the input manuscript")

    heading = document.add_paragraph(style="Heading 1")
    heading.paragraph_format.first_line_indent = Inches(0)
    heading.paragraph_format.page_break_before = True
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    number_run = heading.add_run("CHAPTER THREE")
    number_run.add_break(WD_BREAK.LINE)
    title_run = heading.add_run("THE ORCHESTRATORS")
    set_run_font(title_run, size=18)

    document.add_paragraph(style="Normal")

    for raw_block in parse_blocks(chapter_file):
        block = raw_block
        display = block.startswith("[DISPLAY]")
        readout = block.startswith("[READOUT]")
        centered = block.startswith("[CENTER]")
        if display:
            block = block[len("[DISPLAY]") :].lstrip()
        elif readout:
            block = block[len("[READOUT]") :].lstrip()
        elif centered:
            block = block[len("[CENTER]") :].lstrip()

        paragraph = document.add_paragraph(style="Normal")
        if display or centered:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Inches(0)
        elif readout:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Inches(0)

        add_inline_runs(paragraph, block)
        if display or readout:
            for run in paragraph.runs:
                run.bold = True


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    parser.add_argument("--chapter", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--expected-input-sha256", required=True)
    args = parser.parse_args()

    input_path = Path(args.input)
    expected = args.expected_input_sha256.upper()
    actual = sha256(input_path)
    if actual != expected:
        raise RuntimeError(f"Input manuscript changed: expected {expected}, found {actual}")

    document = Document(input_path)
    add_chapter(document, Path(args.chapter))

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    print(f"{output_path.resolve()}\n{sha256(output_path)}")


if __name__ == "__main__":
    main()
