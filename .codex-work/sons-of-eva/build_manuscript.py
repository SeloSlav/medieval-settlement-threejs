from __future__ import annotations

import argparse
import json
import re
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


INLINE_ITALIC = re.compile(r"(<i>.*?</i>)")


def set_run_font(run, name: str | None = None, size: float | None = None):
    if name:
        run.font.name = name
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:ascii"), name)
        rfonts.set(qn("w:hAnsi"), name)
    if size:
        run.font.size = Pt(size)


def clear_body_keep_final_section(document: Document):
    body = document._element.body
    final_sect_pr = body.sectPr
    for child in list(body):
        if child is not final_sect_pr:
            body.remove(child)


def remove_source_content_relationships(document: Document):
    source_only_types = {
        RT.HYPERLINK,
        RT.IMAGE,
    }
    for relationship_id, relationship in list(document.part.rels.items()):
        if relationship.reltype in source_only_types:
            document.part.drop_rel(relationship_id)


def add_centered_blank(document: Document):
    paragraph = document.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Inches(0)
    return paragraph


def add_inline_runs(paragraph, text: str):
    for part in INLINE_ITALIC.split(text):
        if not part:
            continue
        if part.startswith("<i>") and part.endswith("</i>"):
            run = paragraph.add_run(part[3:-4])
            run.italic = True
        else:
            paragraph.add_run(part)


def parse_chapter(path: Path):
    raw = path.read_text(encoding="utf-8").strip()
    blocks = [block.strip() for block in re.split(r"\r?\n\s*\r?\n", raw) if block.strip()]
    return blocks


def add_title_page(document: Document):
    for _ in range(3):
        add_centered_blank(document)

    series = document.add_paragraph(style="Normal")
    series.alignment = WD_ALIGN_PARAGRAPH.CENTER
    series.paragraph_format.first_line_indent = Inches(0)
    run = series.add_run("БABUSHKA")
    set_run_font(run, "Bookman Old Style", 40)
    run.add_break(WD_BREAK.LINE)

    title = document.add_paragraph(style="Normal")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Inches(0)
    run = title.add_run("SONS OF EVA")
    set_run_font(run, "Bookman Old Style", 22)

    for _ in range(2):
        add_centered_blank(document)

    author = document.add_paragraph(style="Normal")
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.first_line_indent = Inches(0)
    run = author.add_run("By Martin Erlic")
    set_run_font(run, "Bookman Old Style", 16)
    run.add_break(WD_BREAK.PAGE)


def add_chapter(
    document: Document,
    chapter_number: str,
    chapter_title: str,
    blocks: list[str],
    start_new_page: bool = False,
):
    heading = document.add_paragraph(style="Heading 1")
    heading.paragraph_format.first_line_indent = Inches(0)
    heading.paragraph_format.page_break_before = start_new_page
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    number_run = heading.add_run(chapter_number)
    number_run.add_break(WD_BREAK.LINE)
    title_run = heading.add_run(chapter_title)
    set_run_font(title_run, size=18)

    document.add_paragraph(style="Normal")

    for block in blocks:
        display = block.startswith("[DISPLAY]")
        readout = block.startswith("[READOUT]")
        centered = block.startswith("[CENTER]")
        if display:
            block = block[len("[DISPLAY]") :].lstrip()
        elif readout:
            block = block[len("[READOUT]") :].lstrip()
        elif centered:
            block = block[len("[CENTER]") :].lstrip()

        paragraph = document.add_paragraph(style="Normal")
        if display or centered:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Inches(0)
        elif readout:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Inches(0)
        add_inline_runs(paragraph, block)
        if display or readout:
            for run in paragraph.runs:
                run.bold = True


def normalize_single_section(document: Document):
    section = document.sections[-1]
    section.page_width = Inches(6)
    section.page_height = Inches(9)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    section.different_first_page_header_footer = False
    document.settings.odd_and_even_pages_header_footer = False


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--reference", required=True)
    parser.add_argument("--manifest", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()

    document = Document(args.reference)
    clear_body_keep_final_section(document)
    remove_source_content_relationships(document)
    normalize_single_section(document)

    document.core_properties.title = "Sons of Eva"
    document.core_properties.subject = "A Babushka novella"
    document.core_properties.author = "Martin Erlic"
    document.core_properties.keywords = "Babushka, Sons of Eva"
    document.core_properties.comments = ""
    document.core_properties.last_modified_by = ""
    document.core_properties.created = datetime.now(timezone.utc)
    document.core_properties.modified = datetime.now(timezone.utc)

    add_title_page(document)
    manifest_path = Path(args.manifest)
    chapters = json.loads(manifest_path.read_text(encoding="utf-8"))
    for index, chapter in enumerate(chapters):
        chapter_path = manifest_path.parent / chapter["file"]
        add_chapter(
            document,
            chapter["number"],
            chapter["title"],
            parse_chapter(chapter_path),
            start_new_page=index > 0,
        )

    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    print(str(output.resolve()))


if __name__ == "__main__":
    main()
