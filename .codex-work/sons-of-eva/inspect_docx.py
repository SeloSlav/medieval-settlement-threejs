from __future__ import annotations

import argparse
import hashlib
import json
import zipfile
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn


def emu_inches(value):
    return None if value is None else round(value / 914400, 4)


def pt(value):
    return None if value is None else round(value.pt, 3)


def color_of(font):
    try:
        return None if font.color.rgb is None else str(font.color.rgb)
    except Exception:
        return None


def font_names(run):
    rpr = run._element.rPr
    if rpr is None or rpr.rFonts is None:
        return {}
    attrs = {}
    for key in ("ascii", "hAnsi", "eastAsia", "cs", "asciiTheme", "hAnsiTheme"):
        value = rpr.rFonts.get(qn(f"w:{key}"))
        if value:
            attrs[key] = value
    return attrs


def paragraph_record(index, paragraph):
    fmt = paragraph.paragraph_format
    direct_runs = []
    for run_index, run in enumerate(paragraph.runs):
        if not run.text and run._element.xpath('.//w:br'):
            text = "<BREAK>"
        else:
            text = run.text
        direct_runs.append(
            {
                "index": run_index,
                "text": text,
                "bold": run.bold,
                "italic": run.italic,
                "underline": bool(run.underline) if run.underline is not None else None,
                "font_name": run.font.name,
                "font_names": font_names(run),
                "font_size_pt": pt(run.font.size),
                "color": color_of(run.font),
                "small_caps": run.font.small_caps,
                "all_caps": run.font.all_caps,
                "page_break_before_run": bool(run._element.xpath('.//w:br[@w:type="page"]')),
            }
        )
    ppr = paragraph._p.pPr
    sect_pr = bool(ppr is not None and ppr.sectPr is not None)
    return {
        "index": index,
        "text": paragraph.text,
        "style": paragraph.style.name if paragraph.style is not None else None,
        "alignment": None if paragraph.alignment is None else int(paragraph.alignment),
        "left_indent_in": emu_inches(fmt.left_indent),
        "right_indent_in": emu_inches(fmt.right_indent),
        "first_line_indent_in": emu_inches(fmt.first_line_indent),
        "space_before_pt": pt(fmt.space_before),
        "space_after_pt": pt(fmt.space_after),
        "line_spacing": str(fmt.line_spacing) if fmt.line_spacing is not None else None,
        "line_spacing_rule": None if fmt.line_spacing_rule is None else int(fmt.line_spacing_rule),
        "keep_with_next": fmt.keep_with_next,
        "keep_together": fmt.keep_together,
        "page_break_before": fmt.page_break_before,
        "widow_control": fmt.widow_control,
        "section_break_after": sect_pr,
        "runs": direct_runs,
    }


def style_record(style):
    fmt = getattr(style, "paragraph_format", None)
    font = getattr(style, "font", None)
    base_style = getattr(style, "base_style", None)
    next_style = getattr(style, "next_paragraph_style", None) if style.type == 1 else None
    return {
        "style_id": style.style_id,
        "name": style.name,
        "type": int(style.type),
        "base_style": base_style.name if base_style is not None else None,
        "next_style": next_style.name if next_style is not None else None,
        "font_name": None if font is None else font.name,
        "font_size_pt": None if font is None else pt(font.size),
        "font_bold": None if font is None else font.bold,
        "font_italic": None if font is None else font.italic,
        "font_color": None if font is None else color_of(font),
        "paragraph": None if fmt is None else {
            "alignment": None if fmt.alignment is None else int(fmt.alignment),
            "left_indent_in": emu_inches(fmt.left_indent),
            "right_indent_in": emu_inches(fmt.right_indent),
            "first_line_indent_in": emu_inches(fmt.first_line_indent),
            "space_before_pt": pt(fmt.space_before),
            "space_after_pt": pt(fmt.space_after),
            "line_spacing": str(fmt.line_spacing) if fmt.line_spacing is not None else None,
            "line_spacing_rule": None if fmt.line_spacing_rule is None else int(fmt.line_spacing_rule),
            "keep_with_next": fmt.keep_with_next,
            "keep_together": fmt.keep_together,
            "page_break_before": fmt.page_break_before,
            "widow_control": fmt.widow_control,
        },
    }


def section_record(index, section):
    def hf_text(part):
        return [p.text for p in part.paragraphs]

    return {
        "index": index,
        "start_type": int(section.start_type),
        "orientation": "landscape" if section.orientation == WD_ORIENT.LANDSCAPE else "portrait",
        "page_width_in": emu_inches(section.page_width),
        "page_height_in": emu_inches(section.page_height),
        "top_margin_in": emu_inches(section.top_margin),
        "bottom_margin_in": emu_inches(section.bottom_margin),
        "left_margin_in": emu_inches(section.left_margin),
        "right_margin_in": emu_inches(section.right_margin),
        "header_distance_in": emu_inches(section.header_distance),
        "footer_distance_in": emu_inches(section.footer_distance),
        "different_first_page_header_footer": section.different_first_page_header_footer,
        "header_linked": section.header.is_linked_to_previous,
        "footer_linked": section.footer.is_linked_to_previous,
        "header_text": hf_text(section.header),
        "first_page_header_text": hf_text(section.first_page_header),
        "even_page_header_text": hf_text(section.even_page_header),
        "footer_text": hf_text(section.footer),
        "first_page_footer_text": hf_text(section.first_page_footer),
        "even_page_footer_text": hf_text(section.even_page_footer),
    }


def inspect(path: Path):
    doc = Document(path)
    paragraphs = [paragraph_record(i, p) for i, p in enumerate(doc.paragraphs)]
    styles = [style_record(s) for s in doc.styles]
    tables = []
    for t_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([[p.text for p in cell.paragraphs] for cell in row.cells])
        tables.append({"index": t_idx, "style": table.style.name if table.style else None, "rows": rows})
    with zipfile.ZipFile(path) as archive:
        parts = []
        for info in archive.infolist():
            data = archive.read(info.filename)
            parts.append(
                {
                    "path": info.filename,
                    "size": len(data),
                    "sha256": hashlib.sha256(data).hexdigest(),
                }
            )
        doc_xml = archive.read("word/document.xml").decode("utf-8", errors="replace")
        field_tokens = Counter()
        for token in ("TOC", "PAGE", "NUMPAGES", "REF", "PAGEREF"):
            field_tokens[token] = doc_xml.count(token)
    return {
        "path": str(path.resolve()),
        "size": path.stat().st_size,
        "sha256": hashlib.sha256(path.read_bytes()).hexdigest(),
        "core_properties": {
            "title": doc.core_properties.title,
            "subject": doc.core_properties.subject,
            "author": doc.core_properties.author,
            "keywords": doc.core_properties.keywords,
            "comments": doc.core_properties.comments,
            "last_modified_by": doc.core_properties.last_modified_by,
            "created": str(doc.core_properties.created),
            "modified": str(doc.core_properties.modified),
        },
        "settings": {
            "odd_and_even_pages_header_footer": doc.settings.odd_and_even_pages_header_footer,
        },
        "sections": [section_record(i + 1, s) for i, s in enumerate(doc.sections)],
        "styles": styles,
        "paragraphs": paragraphs,
        "tables": tables,
        "inline_shapes": len(doc.inline_shapes),
        "package_parts": parts,
        "field_token_counts": field_tokens,
        "style_counts": Counter(p["style"] for p in paragraphs),
    }


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("input")
    parser.add_argument("--json", required=True)
    parser.add_argument("--text", required=True)
    args = parser.parse_args()
    path = Path(args.input)
    report = inspect(path)
    Path(args.json).write_text(json.dumps(report, indent=2, ensure_ascii=False, default=dict), encoding="utf-8")
    lines = []
    for p in report["paragraphs"]:
        if p["text"] or p["style"].startswith("Heading") or p["section_break_after"]:
            lines.append(f'[{p["index"]:05d}] [{p["style"]}] {p["text"]}')
    for table in report["tables"]:
        lines.append(f'[[TABLE {table["index"]} style={table["style"]}]]')
        for row in table["rows"]:
            lines.append(" | ".join(" / ".join(cell) for cell in row))
    Path(args.text).write_text("\n".join(lines), encoding="utf-8")


if __name__ == "__main__":
    main()
